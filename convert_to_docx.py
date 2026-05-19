"""Convert Documentation.md to Documentation.docx using python-docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = "Documentation.md"
DOCX_FILE = "Documentation.docx"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pb.append(bottom)
    pPr.append(pb)
    return p


def style_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x3A, 0x6A)
    return h


def parse_inline(run_text, paragraph):
    """Parse inline bold/code markers and add runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", run_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def parse_table(doc, lines):
    """Parse a markdown table block and add a docx table."""
    rows = []
    for line in lines:
        if re.match(r"\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= cols:
                break
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            if i == 0:
                set_cell_bg(cell, "1A3A6E")
                run = p.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
            else:
                parse_inline(cell_text, p)
                for run in p.runs:
                    run.font.size = Pt(9)
                if i % 2 == 0:
                    set_cell_bg(cell, "F0F4F8")

    doc.add_paragraph()


def convert(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        # --- Code block ---
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                code_para = doc.add_paragraph()
                code_para.style = "No Spacing"
                for cl in code_lines:
                    r = code_para.add_run(cl + "\n")
                    r.font.name = "Courier New"
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                pf = code_para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F4F4F4")
                pf.append(shd)
                doc.add_paragraph()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Table detection ---
        if "|" in line and not line.strip().startswith("#"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                parse_table(doc, table_lines)
                table_lines = []
                in_table = False

        # --- Horizontal rule ---
        if re.match(r"^---+\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # --- Headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip bold markers from heading text
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            style_heading(doc, text, level)
            i += 1
            continue

        # --- Blockquote ---
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.style = "No Spacing"
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")
            pPr.append(ind)
            pb = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "8")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "4472C4")
            pb.append(left)
            pPr.append(pb)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
            i += 1
            continue

        # --- Bullet list ---
        m = re.match(r"^(\s*)[•\-\*]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.5 * (indent // 2 + 1))
            parse_inline(text, p)
            i += 1
            continue

        # --- Numbered list ---
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            text = m.group(1).strip()
            p = doc.add_paragraph(style="List Number")
            parse_inline(text, p)
            i += 1
            continue

        # --- Empty line ---
        if line.strip() == "":
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        parse_inline(line, p)
        i += 1

    # Flush any trailing table
    if in_table and table_lines:
        parse_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    convert(
        os.path.join(base, MD_FILE),
        os.path.join(base, DOCX_FILE),
    )
